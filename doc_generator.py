from docx import Document
from io import BytesIO


def create_doc(title, datetime_str, participants, mom):
    """
    Build the MOM .docx. Defensive against missing keys, None values,
    and empty sections so a slightly malformed MOM never crashes the
    download path.
    """
    mom = mom or {}

    doc = Document()

    doc.add_heading("Minutes of Meeting", 0)

    doc.add_heading("Meeting Details", 1)
    doc.add_paragraph(f"Title: {title or 'Untitled'}")
    doc.add_paragraph(f"Date & Time: {datetime_str or 'N/A'}")

    doc.add_paragraph("Participants:")
    for p in str(participants or "").split(","):
        if p.strip():
            doc.add_paragraph(p.strip(), style="List Bullet")

    doc.add_heading("Summary", 1)
    doc.add_paragraph(str(mom.get("summary", "") or "No summary provided."))

    doc.add_heading("Key Decisions", 1)
    decisions = [str(d).strip() for d in mom.get("decisions", []) if str(d).strip()]
    if decisions:
        for d in decisions:
            doc.add_paragraph(d, style="List Bullet")
    else:
        doc.add_paragraph("No decisions recorded.")

    doc.add_heading("Risks", 1)
    risks = [str(r).strip() for r in mom.get("risks", []) if str(r).strip()]
    if risks:
        for r in risks:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph("No risks recorded.")

    doc.add_heading("Action Items", 1)
    actions = mom.get("actions", []) or []

    if actions:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = "Task"
        header[1].text = "Owner"
        header[2].text = "Deadline"

        for a in actions:
            if not isinstance(a, dict):
                # Tolerate stray string entries instead of crashing
                a = {"task": str(a)}
            row = table.add_row().cells
            row[0].text = str(a.get("task", "") or "")
            row[1].text = str(a.get("owner", "") or "TBD")
            row[2].text = str(a.get("deadline", "") or "TBD")
    else:
        doc.add_paragraph("No action items recorded.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
